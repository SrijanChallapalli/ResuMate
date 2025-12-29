"""
File extraction utilities for PDF, DOCX, and TXT files.
"""
import os
import re
from typing import Tuple, Optional
from pathlib import Path


class FileExtractor:
    """Extracts text from various file formats."""
    
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
    MAX_EXTRACTED_TEXT = 25000  # 25k chars max
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    
    @staticmethod
    def validate_file(filename: str, file_size: int) -> Tuple[bool, Optional[str]]:
        """
        Validate file type and size.
        Returns: (is_valid, error_message)
        """
        # Handle None or empty filename
        if not filename:
            return False, "Filename is required"
        
        # Check file extension
        try:
            file_ext = Path(filename).suffix.lower()
        except (TypeError, AttributeError) as e:
            return False, f"Invalid filename format: {str(e)}"
        
        if not file_ext or file_ext not in FileExtractor.ALLOWED_EXTENSIONS:
            return False, f"Unsupported file type. Allowed: {', '.join(FileExtractor.ALLOWED_EXTENSIONS)}"
        
        # Check file size
        if file_size > FileExtractor.MAX_FILE_SIZE:
            return False, f"File too large. Maximum size: {FileExtractor.MAX_FILE_SIZE / (1024*1024):.1f}MB"
        
        if file_size == 0:
            return False, "File is empty"
        
        return True, None
    
    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Sanitize filename to prevent path traversal and other issues."""
        # Remove path components
        filename = os.path.basename(filename)
        # Remove dangerous characters
        filename = re.sub(r'[<>:"|?*\x00-\x1f]', '', filename)
        # Limit length
        if len(filename) > 255:
            filename = filename[:255]
        return filename
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            from pypdf import PdfReader
            import io
            
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception:
                    continue
            
            extracted = '\n\n'.join(text_parts)
            # Truncate to max length
            if len(extracted) > FileExtractor.MAX_EXTRACTED_TEXT:
                extracted = extracted[:FileExtractor.MAX_EXTRACTED_TEXT]
            
            return extracted.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            import io
            
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            extracted = '\n\n'.join(text_parts)
            # Truncate to max length
            if len(extracted) > FileExtractor.MAX_EXTRACTED_TEXT:
                extracted = extracted[:FileExtractor.MAX_EXTRACTED_TEXT]
            
            return extracted.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from TXT file."""
        try:
            # Try UTF-8 first
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1 (covers most cases)
                text = file_content.decode('latin-1', errors='ignore')
            
            # Truncate to max length
            if len(text) > FileExtractor.MAX_EXTRACTED_TEXT:
                text = text[:FileExtractor.MAX_EXTRACTED_TEXT]
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")
    
    @staticmethod
    def extract_text(file_content: bytes, filename: str) -> Tuple[str, dict]:
        """
        Extract text from file based on extension.
        Returns: (extracted_text, metadata)
        """
        if not filename:
            raise ValueError("Filename is required for text extraction")
        
        try:
            file_ext = Path(filename).suffix.lower()
        except (TypeError, AttributeError):
            raise ValueError(f"Invalid filename format: {filename}")
        
        if not file_ext:
            raise ValueError(f"Cannot determine file type from filename: {filename}")
        
        pages = None
        
        if file_ext == '.pdf':
            text = FileExtractor.extract_text_from_pdf(file_content)
            # Try to estimate pages (rough estimate)
            pages = max(1, len(text) // 2000)  # Rough estimate: 2000 chars per page
        elif file_ext == '.docx':
            text = FileExtractor.extract_text_from_docx(file_content)
        elif file_ext == '.txt':
            text = FileExtractor.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        if not text or len(text.strip()) < 10:
            raise ValueError("Extracted text is too short or empty. Please ensure the file contains readable text.")
        
        meta = {
            "filename": FileExtractor.sanitize_filename(filename),
            "pages": pages,
            "textLength": len(text)
        }
        
        return text, meta

